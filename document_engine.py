import pandas as pd
from docx import Document
import io

def replace_text_in_paragraphs(paragraphs, state_dict):
    """Iterates through paragraphs and replaces string placeholders."""
    for paragraph in paragraphs:
        for key, value in state_dict.items():
            # Only process scalar values (strings, numbers), ignore DataFrames
            if not isinstance(value, pd.DataFrame) and value is not None:
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))

def insert_table_after_paragraph(paragraph, df):
    """Builds a Word table from a DataFrame and inserts it immediately after the specified paragraph."""
    document = paragraph._parent
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    
    # Add Headers
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = str(col_name)
        
    # Add Data Rows
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    # Move table XML right after the paragraph XML
    paragraph._p.addnext(table._tbl)

def generate_master_document(template_path, case_state):
    """
    Main generation function combining scalar replacements and dynamic tables.
    Returns a bytes object for Streamlit downloading.
    """
    doc = Document(template_path)
    
    # 1. Replace scalar text in main body paragraphs
    replace_text_in_paragraphs(doc.paragraphs, case_state)
    
    # 2. Replace scalar text in existing tables (e.g., the Index table in LOC filing)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(cell.paragraphs, case_state)

    # 3. Dynamic Table Injection for DataFrames
    table_placeholders = {
        "{{df_creditors}}": case_state.get("df_creditors"),
        "{{df_suspended_mgmt}}": case_state.get("df_suspended_mgmt"),
        "{{df_expenses}}": case_state.get("df_expenses")
    }

    for paragraph in doc.paragraphs:
        for placeholder, df in table_placeholders.items():
            if placeholder in paragraph.text:
                if df is not None and not df.empty:
                    insert_table_after_paragraph(paragraph, df)
                # Clear the placeholder text regardless of whether a table was generated
                paragraph.text = paragraph.text.replace(placeholder, "")

    # Save to a virtual file for download in the browser
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()